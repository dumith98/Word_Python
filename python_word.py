from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc =  Document()

titulo = doc.add_heading(level = 0)
titulo.add_run('Word com ').bold = True
titulo.add_run('Python').italic = True

doc.add_paragraph('''Pode parecer estranho usar o python para editar um documento word,
 mas acredite ou não tem as suas vantagens. Pra quem ja esta acostuamdo a usar o VSCode,
 editar um texto desta maneira pode ser mais natural.''')
doc.add_paragraph('Somos capazes de facilmente fazer de tudo no word com Python, como: ')

doc.add_paragraph('Colocar imagens;', style = 'List Number')
doc.add_paragraph('Estilizar fontes; ', style = 'List Number')
doc.add_paragraph('Mudar alinhamentos;', style = 'List Number')
doc.add_paragraph('Inserir tabelas;', style = 'List Number')
doc.add_paragraph('Fazer uma lista;', style = 'List Number')
doc.add_paragraph('e muito mais;', style = 'List Number')

doc.add_picture('fotocachorro.jpg', width = Inches(2))
estilo2 = doc.styles.add_style('Legenda', WD_STYLE_TYPE.PARAGRAPH)
estilo2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_HI
fonte2 = estilo2.font
fonte2.size = Pt(7)
fonte2.italic = True
legenda = doc.add_paragraph()
legenda.style = doc.styles['Legenda']
legenda.add_run('Cachorro fofo')


estilo1 = doc.styles.add_style('Estilizado', WD_STYLE_TYPE.PARAGRAPH)
estilo1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
fonte = estilo1.font
fonte.name = 'Castellar'
fonte.size = Pt(20)
fonte.color.rgb = RGBColor(0, 0, 255)
tabela_titulo = doc.add_paragraph()
tabela_titulo.style = doc.styles['Estilizado']
tabela_titulo.add_run('Eis uma Tabela.')

dados_tabela = [['Gato', 'Docil', 'Casa'], ['Cachorro', 'Docil', 'Casa'], ['Leão', 'Agressivo', 'Selva'], ['Onça', 'Agressivo', 'Floresta'], ['Lobo', 'Agressivo', 'Tundra']]

tabela = doc.add_table(rows = 1, cols = 3)
celula_titulo = tabela.rows[0].cells
celula_titulo[0].text = 'Nome'
celula_titulo[1].text = 'Sociabilidade'
celula_titulo[2].text = 'Local'

for nome, social, lugar in dados_tabela:
    celula_nova_linha = tabela.add_row().cells
    celula_nova_linha[0].text = nome
    celula_nova_linha[1].text = social
    celula_nova_linha[2].text = lugar

meio = doc.add_heading(level = 1)
meio.add_run('Vale lembrar...').italic = True
doc.add_paragraph(' Concordo que hoje em dia se usa-se mais PDF e planilhas Excel para se fazer de tudo e o Word pode parecer um tanto ultrapassado para projetos profissionais, mas nunca se sabe...', 'Quote')
doc.add_paragraph('É sempre bom estar preparado para uma situação que precise usar o Word, afinal ele é uma boa ferramenta para documentos padronizados.', 'Intense Quote')


doc.save('desafio_word.docx')